# -*- coding: utf-8 -*-
"""
Build 5 Word (.docx) lab reports from the Markdown sources in docs/reports/.

Rules:
- Cover page (placeholders kept) + numbered Word headings.
- Markdown tables  -> real Word tables (Table Grid, bold header).
- Lab 1-4 SQL code -> italic monospace (Consolas) text blocks.
- Lab 5 SQL code   -> syntax-highlighted PNG image (light background).
- Lab 5 result blocks -> monospace text blocks.
- Diagrams (existing PNGs) inserted per template with numbered captions.
- No emojis / checkbox glyphs / ">" callouts / marketing labels.

Run:  python docs/reports/_build_lab_docx.py
"""
import io
import os
import re
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PIL import Image
from pygments import highlight
from pygments.lexers import SqlLexer
from pygments.formatters import ImageFormatter

HERE = os.path.dirname(os.path.abspath(__file__))
DOCS = os.path.dirname(HERE)            # docs/
REPO = os.path.dirname(DOCS)            # repo root

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
BODY_SIZE = 13
TABLE_SIZE = 11
LINE_SPACING = 1.4
PAGE_WIDTH_IN = 6.3                     # usable width inside A4 margins

# ----------------------------------------------------------------------------
# Per-lab figure plan (decided from the teacher .docx templates in docs/).
#   key = source markdown filename
#   "title" -> document title shown on the cover
#   "figures" = list of extra figures to inject that are NOT already in the .md
#       each: (anchor_heading_prefix, position, image, caption)
#       position: "after" the matched heading block.
# erd.png is already referenced inside lab2/lab4 .md, so it is handled by the
# normal image parser; only the additional design diagrams are injected here.
# ----------------------------------------------------------------------------
FIG_PLAN = {
    "lab4_relational_design_process.md": [
        ("## 6. Physical Design", "docs/block_diagram.png",
         "System architecture overview of the LMS database and its programmable objects."),
        ("### 7.2.", "docs/flowchart_submission.png",
         "Submission process flow enforced by sp_SubmitAssignment and trg_Submissions_Policy."),
    ],
}

# A standalone italic caption line already written in the markdown, e.g.
# "*Figure 1. Entity-Relationship Diagram ...*" — skipped because the builder
# renders its own numbered caption under every figure.
CAPTION_RE = re.compile(r"^\*?\s*Figure\s+\d+\.", re.IGNORECASE)

LAB_FILES = [
    "lab1_data_models.md",
    "lab2_entities_fds_keys.md",
    "lab3_anomalies_and_normalization.md",
    "lab4_relational_design_process.md",
    "lab5_sql_programming.md",
]

SQL_INLINE_RE = re.compile(r"^(fn_|sp_|trg_|vw_|IX_|CK_|FK_|UQ_|DF_|PK_)")
# A backticked token that is a file/path reference rather than a SQL object.
PATHLIKE_RE = re.compile(r"[\\/]|\.(sql|md|png|txt|py|docx|pdf|csv|bat)\b", re.IGNORECASE)


def is_sql_object(token):
    """Treat an inline `code` span as a SQL object name (-> italic) unless it
    is a file path / file name reference."""
    t = token.strip()
    if not t or PATHLIKE_RE.search(t):
        return False
    return True


# ===========================================================================
# Low-level docx helpers
# ===========================================================================
def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_normal(doc):
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(BODY_SIZE)
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    pf = st.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)

    for name, size in (("Heading 1", 15), ("Heading 2", 13.5), ("Heading 3", 13)):
        h = doc.styles[name]
        h.font.name = BODY_FONT
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.font.italic = False


def add_heading_para(doc, text, level):
    """Render a section heading the way the teacher templates do: a plain
    Normal-style paragraph, bold, LEFT-aligned and flush to the left margin
    (no centering, no heading-style indentation). Sizes step down by level."""
    sizes = {2: 14, 3: 13, 4: 12.5}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.left_indent = Inches(0)
    pf.first_line_indent = Inches(0)
    pf.space_before = Pt(10 if level == 2 else 8)
    pf.space_after = Pt(4)
    pf.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(sizes.get(level, 12.5))
    r.font.color.rgb = RGBColor(0, 0, 0)
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    return p


def add_runs(paragraph, text, base_italic=False, base_mono=False, sql_inline=False):
    """Render inline markdown (**bold**, *italic*, `code`, [t](u)) into runs."""
    token = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]*?\))")
    pos = 0
    for m in token.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos:m.start()], base_italic, base_mono)
        tok = m.group(0)
        if tok.startswith("**"):
            _run(paragraph, tok[2:-2], base_italic, base_mono, bold=True)
        elif tok.startswith("`"):
            inner = tok[1:-1]
            ital = (base_italic or is_sql_object(inner)) if sql_inline else base_italic
            _run(paragraph, inner, italic=ital, mono=True)
        elif tok.startswith("["):
            label = re.match(r"\[([^\]]+?)\]", tok).group(1)
            _run(paragraph, label, base_italic, base_mono)
        else:  # *italic*
            _run(paragraph, tok[1:-1], italic=True, mono=base_mono)
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:], base_italic, base_mono)


def _run(paragraph, text, italic=False, mono=False, bold=False):
    if text == "":
        return
    r = paragraph.add_run(text)
    r.italic = italic
    r.bold = bold
    if mono:
        r.font.name = MONO_FONT
        rpr = r._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), MONO_FONT)
        rfonts.set(qn("w:hAnsi"), MONO_FONT)


# ===========================================================================
# Markdown -> block tokens
# ===========================================================================
def parse_blocks(md_lines):
    """Return a list of (kind, payload) blocks."""
    blocks = []
    i, n = 0, len(md_lines)
    while i < n:
        line = md_lines[i]
        stripped = line.strip()

        # fenced code
        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            i += 1
            buf = []
            while i < n and not md_lines[i].strip().startswith("```"):
                buf.append(md_lines[i].rstrip("\n"))
                i += 1
            i += 1  # closing fence
            blocks.append(("code", {"lang": lang, "lines": buf}))
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            blocks.append(("heading", {"level": len(m.group(1)), "text": m.group(2).strip()}))
            i += 1
            continue

        # image  ![alt](path)
        m = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", stripped)
        if m:
            blocks.append(("image", {"alt": m.group(1), "path": m.group(2)}))
            i += 1
            continue

        # table
        if "|" in stripped and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", md_lines[i + 1]) \
           and "-" in md_lines[i + 1]:
            tbl = []
            while i < n and "|" in md_lines[i] and md_lines[i].strip():
                tbl.append(md_lines[i].strip())
                i += 1
            # drop separator row (index 1)
            rows = [r for j, r in enumerate(tbl) if j != 1]
            parsed = []
            for r in rows:
                cells = r.strip().strip("|").split("|")
                parsed.append([c.strip() for c in cells])
            blocks.append(("table", parsed))
            continue

        # blockquote
        if stripped.startswith(">"):
            buf = []
            while i < n and md_lines[i].strip().startswith(">"):
                buf.append(re.sub(r"^\s*>\s?", "", md_lines[i].rstrip("\n")))
                i += 1
            blocks.append(("quote", buf))
            continue

        # horizontal rule
        if re.match(r"^---+$", stripped):
            i += 1
            continue

        # list item
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            ordered = bool(re.match(r"\d+\.", m.group(2)))
            blocks.append(("li", {"indent": indent, "ordered": ordered, "text": m.group(3).strip()}))
            i += 1
            continue

        # blank
        if stripped == "":
            i += 1
            continue

        # paragraph (gather continuation lines)
        buf = [stripped]
        i += 1
        while i < n:
            nxt = md_lines[i].strip()
            if (nxt == "" or nxt.startswith("#") or nxt.startswith("```")
                    or nxt.startswith(">") or nxt.startswith("![")
                    or re.match(r"^(\s*)([-*]|\d+\.)\s+", md_lines[i])
                    or ("|" in nxt and i + 1 < n and "-" in md_lines[i + 1] and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", md_lines[i + 1]))
                    or re.match(r"^---+$", nxt)):
                break
            buf.append(nxt)
            i += 1
        blocks.append(("para", " ".join(buf)))
    return blocks


# ===========================================================================
# Cover page
# ===========================================================================
def build_cover(doc, title, meta):
    def centered(txt, size, bold=False, italic=False, after=6, mono=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(txt)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = BODY_FONT
        return p

    for _ in range(2):
        doc.add_paragraph()
    centered(meta.get("Course", "DBI202 — Database Systems").upper(), 16, bold=True, after=4)
    centered("Laboratory Report", 13, after=2)
    for _ in range(3):
        doc.add_paragraph()
    centered(title, 22, bold=True, after=8)
    centered(meta.get("Project", "Online Learning Management System (LMS)"), 14, italic=True, after=4)
    for _ in range(4):
        doc.add_paragraph()

    centered("Group: " + meta.get("Group", "[GROUP NAME]"), 13, bold=True, after=6)
    members = [m.strip() for m in meta.get("Members", "").split(",") if m.strip()]
    centered("Members", 13, bold=True, after=2)
    for mb in members:
        centered(mb, 13, after=2)
    doc.add_paragraph()
    centered("Class: " + meta.get("Class", "[CLASS CODE]"), 13, after=2)
    centered("Date: " + meta.get("Date", "[SUBMISSION DATE]"), 13, after=2)

    doc.add_page_break()


# ===========================================================================
# Code-as-image (Lab 5)
# ===========================================================================
def code_to_png(code):
    fmt = ImageFormatter(style="friendly", font_name=MONO_FONT, font_size=26,
                         line_numbers=False, image_pad=16)
    data = highlight(code, SqlLexer(), fmt)
    return data


def add_code_image(doc, code):
    data = code_to_png(code)
    bio = io.BytesIO(data)
    im = Image.open(bio)
    w_px = im.size[0]
    bio.seek(0)
    width_in = min(PAGE_WIDTH_IN, w_px / 96.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(bio, width=Inches(width_in))


# ===========================================================================
# Monospace text block (plain code / results)
# ===========================================================================
def add_mono_block(doc, lines, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    text = "\n".join(lines)
    r = p.add_run(text)
    r.font.name = MONO_FONT
    r.font.size = Pt(10.5)
    r.italic = italic
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), MONO_FONT)
    rfonts.set(qn("w:hAnsi"), MONO_FONT)
    # light shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


# ===========================================================================
# Table
# ===========================================================================
def add_table(doc, rows):
    ncol = max(len(r) for r in rows)
    tbl = doc.add_table(rows=0, cols=ncol)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ridx, row in enumerate(rows):
        cells = tbl.add_row().cells
        for cidx in range(ncol):
            txt = row[cidx] if cidx < len(row) else ""
            cell = cells[cidx]
            cell.paragraphs[0].text = ""
            par = cell.paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            par.paragraph_format.line_spacing = 1.0
            add_runs(par, txt, sql_inline=True)
            for run in par.runs:
                if run.font.size is None:
                    run.font.size = Pt(TABLE_SIZE)
                if ridx == 0:
                    run.bold = True
            if ridx == 0:
                set_cell_background(cell, "D9D9D9")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ===========================================================================
# Figures
# ===========================================================================
def add_figure(doc, img_rel_or_abs, caption, counter):
    path = img_rel_or_abs
    if not os.path.isabs(path):
        # paths in md are like ../erd.png (relative to docs/reports)
        cand = os.path.normpath(os.path.join(HERE, path))
        if not os.path.exists(cand):
            cand = os.path.normpath(os.path.join(REPO, path))
        path = cand
    if not os.path.exists(path):
        print("  [WARN] image not found:", path)
        return counter
    im = Image.open(path)
    w_px = im.size[0]
    width_in = min(PAGE_WIDTH_IN, w_px / 96.0)
    if width_in < 4.0:
        width_in = min(PAGE_WIDTH_IN, 5.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    counter += 1
    r = cap.add_run("Figure %d. %s" % (counter, caption))
    r.italic = True
    r.font.size = Pt(11)
    cap.paragraph_format.space_after = Pt(10)
    return counter


# ===========================================================================
# Build one document
# ===========================================================================
def build_doc(md_path, out_path):
    fname = os.path.basename(md_path)
    lab5 = fname.startswith("lab5")
    with open(md_path, "r", encoding="utf-8") as f:
        md_lines = f.read().split("\n")

    blocks = parse_blocks(md_lines)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)
    style_normal(doc)

    # ---- cover: first heading (#) + first quote block (metadata) ----
    title = None
    meta = {}
    idx = 0
    if blocks and blocks[0][0] == "heading" and blocks[0][1]["level"] == 1:
        title = blocks[0][1]["text"]
        idx = 1
    if idx < len(blocks) and blocks[idx][0] == "quote":
        for ln in blocks[idx][1]:
            mm = re.match(r"\*\*(.+?):\*\*\s*(.*)", ln.strip())
            if mm:
                meta[mm.group(1).strip()] = mm.group(2).strip()
        idx += 1
    build_cover(doc, title or fname, meta)

    fig_counter = 0
    extra_figs = FIG_PLAN.get(fname, [])
    last_heading_text = ""

    for kind, payload in blocks[idx:]:
        if kind == "heading":
            lvl = payload["level"]
            text = payload["text"]
            last_heading_text = text
            add_heading_para(doc, text, lvl)
            # inject planned design figures anchored after the matched heading
            for anchor, img, cap in list(extra_figs):
                a = anchor.lstrip("#").strip()
                if text.strip().startswith(a) or a.startswith(text.strip()):
                    fig_counter = add_figure(doc, img, cap, fig_counter)
                    extra_figs.remove((anchor, img, cap))

        elif kind == "para":
            if CAPTION_RE.match(payload.strip()):
                continue  # md's own figure caption -> builder adds a numbered one
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(p, payload, sql_inline=True)

        elif kind == "quote":
            for ln in payload:
                if not ln.strip():
                    continue
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.left_indent = Inches(0.3)
                add_runs(p, ln, sql_inline=True)

        elif kind == "li":
            ordered = payload["ordered"]
            style = "List Number" if ordered else "List Bullet"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(p, payload["text"], sql_inline=True)

        elif kind == "table":
            add_table(doc, payload)

        elif kind == "image":
            cap = payload["alt"] or "Diagram"
            fig_counter = add_figure(doc, payload["path"], cap, fig_counter)

        elif kind == "code":
            lang = payload["lang"]
            lines = payload["lines"]
            if not lines:
                continue
            if lab5 and lang == "sql":
                add_code_image(doc, "\n".join(lines))
            elif lab5:
                add_mono_block(doc, lines, italic=False)
            else:
                # Lab 1-4: SQL italic monospace; other fenced text monospace plain
                add_mono_block(doc, lines, italic=(lang == "sql"))

    doc.save(out_path)
    return os.path.getsize(out_path)


def main():
    results = []
    for fname in LAB_FILES:
        md_path = os.path.join(HERE, fname)
        if not os.path.exists(md_path):
            print("MISSING:", md_path)
            continue
        out_path = os.path.join(HERE, fname.replace(".md", ".docx"))
        try:
            size = build_doc(md_path, out_path)
        except PermissionError:
            print("PERMISSION ERROR (file open?):", out_path)
            sys.exit(2)
        results.append((os.path.basename(out_path), size))
        print("BUILT %-40s %8.1f KB" % (os.path.basename(out_path), size / 1024.0))
    print("\nDone:", len(results), "files.")


if __name__ == "__main__":
    main()
